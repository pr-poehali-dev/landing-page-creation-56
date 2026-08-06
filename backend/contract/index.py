import json
import os
import io
import time
import psycopg2
import boto3
from docx import Document

MONTHS_RU = ["", "января", "февраля", "марта", "апреля", "мая", "июня",
             "июля", "августа", "сентября", "октября", "ноября", "декабря"]

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "templates", "dogovor_template.docx")


def parse_date(s):
    if not s:
        return None
    y, m, d = s.split("-")[:3]
    d = d[:2]
    return (int(y), int(m), int(d))


def fmt_date(dt):
    y, m, d = dt
    return f"{d:02d}.{m:02d}.{y}"


def build_requisites_text(lead):
    lines = []
    company = lead.get('company') or lead.get('name') or ''
    lines.append(company)
    if lead.get('legal_address'):
        lines.append(f"Юридический адрес: {lead['legal_address']}")
    id_line = []
    if lead.get('inn'):
        id_line.append(f"ИНН {lead['inn']}")
    if lead.get('kpp'):
        id_line.append(f"КПП {lead['kpp']}")
    if id_line:
        lines.append(' '.join(id_line))
    if lead.get('ogrn'):
        lines.append(f"ОГРН {lead['ogrn']}")
    if lead.get('bank_account'):
        lines.append(f"р/с {lead['bank_account']}")
    if lead.get('bank_name'):
        lines.append(f"Банк: {lead['bank_name']}")
    if lead.get('bank_corr_account'):
        lines.append(f"к/с {lead['bank_corr_account']}")
    if lead.get('bank_bik'):
        lines.append(f"БИК {lead['bank_bik']}")
    return '\n'.join(lines)


def fill_contract(lead, contract_no, sign_dt):
    doc = Document(TEMPLATE_PATH)
    p = doc.paragraphs

    day = f"{sign_dt[2]:02d}"
    month = MONTHS_RU[sign_dt[1]]
    year2 = f"{sign_dt[0] % 100:02d}"

    customer_name = lead.get('company') or lead.get('name') or ''
    signer_name = lead.get('signer_name') or lead.get('name') or ''

    start_dt = parse_date(lead.get('start_date'))
    end_dt = parse_date(lead.get('end_date'))

    duration = lead.get('duration') or 0
    days = lead.get('days') or 0
    placement_amount = lead.get('placement_amount') or lead.get('total_price') or 0
    outputs_per_day = 204
    outputs_total = outputs_per_day * days if days else 0
    monthly_amount = round(placement_amount / days * 30) if days else placement_amount

    # --- ДОГОВОР шапка ---
    p[0].runs[2].text = str(contract_no)
    p[2].runs[6].text = day
    p[2].runs[9].text = month
    p[2].runs[12].text = year2
    p[4].runs[7].text = customer_name
    if end_dt:
        p[49].runs[4].text = fmt_date(end_dt)

    # --- Приложение №1 ---
    p[70].runs[2].text = str(contract_no)
    p[70].runs[4].text = fmt_date(sign_dt)
    p[72].runs[5].text = day
    p[72].runs[7].text = month + ' '
    p[72].runs[9].text = year2
    p[74].runs[7].text = customer_name
    p[76].runs[10].text = f"{duration} "
    p[79].runs[6].text = f"{placement_amount:,}".replace(",", " ")

    # --- Таблица реквизитов ---
    t0 = doc.tables[0]
    cell = t0.rows[1].cells[1]
    cell.paragraphs[0].text = ""
    run = cell.paragraphs[0].add_run(build_requisites_text(lead))
    run.font.size = t0.rows[1].cells[0].paragraphs[0].runs[0].font.size
    run.font.name = "Times New Roman"

    sign_para = t0.rows[2].cells[1].paragraphs[1]
    if len(sign_para.runs) >= 5:
        sign_para.runs[4].text = signer_name

    # --- Таблица графика размещения ---
    t1 = doc.tables[1]
    row = t1.rows[1]
    period_text = f"с {fmt_date(start_dt)} по {fmt_date(end_dt)}" if start_dt and end_dt else ""
    row.cells[3].text = period_text
    row.cells[4].text = str(outputs_per_day)
    row.cells[5].text = str(outputs_total)
    row.cells[6].text = f"{monthly_amount:,}".replace(",", " ")

    # --- Подписи в приложении ---
    t2 = doc.tables[2]
    sp = t2.rows[1].cells[1].paragraphs[3]
    if len(sp.runs) >= 4:
        sp.runs[3].text = signer_name

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def handler(event: dict, context) -> dict:
    """Формирует договор в формате .docx на основе реквизитов клиента из CRM и загружает его в файловое хранилище"""
    method = event.get('httpMethod', 'GET')

    cors_headers = {
        'Access-Control-Allow-Origin': '*',
        'Access-Control-Allow-Methods': 'POST, OPTIONS',
        'Access-Control-Allow-Headers': 'Content-Type, X-Admin-Key',
        'Access-Control-Max-Age': '86400'
    }

    if method == 'OPTIONS':
        return {'statusCode': 200, 'headers': cors_headers, 'body': ''}

    if method != 'POST':
        return {
            'statusCode': 405,
            'headers': {**cors_headers, 'Content-Type': 'application/json'},
            'body': json.dumps({'error': 'Метод не поддерживается'}, ensure_ascii=False),
            'isBase64Encoded': False
        }

    admin_key = os.environ.get('ADMIN_KEY', '')
    headers = event.get('headers', {})
    provided = headers.get('X-Admin-Key') or headers.get('x-admin-key', '')

    if admin_key and provided != admin_key:
        return {
            'statusCode': 403,
            'headers': {**cors_headers, 'Content-Type': 'application/json'},
            'body': json.dumps({'error': 'Доступ запрещён'}, ensure_ascii=False),
            'isBase64Encoded': False
        }

    body = json.loads(event.get('body', '{}'))
    lead_id = body.get('leadId')

    if not lead_id:
        return {
            'statusCode': 400,
            'headers': {**cors_headers, 'Content-Type': 'application/json'},
            'body': json.dumps({'error': 'leadId обязателен'}, ensure_ascii=False),
            'isBase64Encoded': False
        }

    dsn = os.environ.get('DATABASE_URL')
    conn = psycopg2.connect(dsn)
    cur = conn.cursor()
    cur.execute(
        "SELECT name, company, duration, days, total_price, start_date, end_date, placement_amount, "
        "inn, kpp, ogrn, legal_address, bank_name, bank_account, bank_bik, bank_corr_account, "
        "signer_name, signer_position "
        f"FROM leads WHERE id = {int(lead_id)}"
    )
    row = cur.fetchone()
    cur.close()
    conn.close()

    if not row:
        return {
            'statusCode': 404,
            'headers': {**cors_headers, 'Content-Type': 'application/json'},
            'body': json.dumps({'error': 'Заявка не найдена'}, ensure_ascii=False),
            'isBase64Encoded': False
        }

    lead = {
        'name': row[0], 'company': row[1], 'duration': row[2], 'days': row[3],
        'total_price': row[4],
        'start_date': row[5].isoformat() if row[5] else None,
        'end_date': row[6].isoformat() if row[6] else None,
        'placement_amount': row[7],
        'inn': row[8], 'kpp': row[9], 'ogrn': row[10], 'legal_address': row[11],
        'bank_name': row[12], 'bank_account': row[13], 'bank_bik': row[14], 'bank_corr_account': row[15],
        'signer_name': row[16], 'signer_position': row[17]
    }

    if not lead.get('inn') and not lead.get('legal_address'):
        return {
            'statusCode': 400,
            'headers': {**cors_headers, 'Content-Type': 'application/json'},
            'body': json.dumps({'error': 'У заявки не заполнены реквизиты клиента'}, ensure_ascii=False),
            'isBase64Encoded': False
        }

    contract_no = body.get('contractNo') or str(lead_id)
    sign_date_str = body.get('signDate')
    sign_dt = parse_date(sign_date_str) if sign_date_str else time.localtime()[0:3]

    docx_bytes = fill_contract(lead, contract_no, sign_dt)

    s3 = boto3.client(
        's3',
        endpoint_url='https://bucket.poehali.dev',
        aws_access_key_id=os.environ['AWS_ACCESS_KEY_ID'],
        aws_secret_access_key=os.environ['AWS_SECRET_ACCESS_KEY']
    )
    file_key = f"contracts/dogovor_{lead_id}_{int(time.time())}.docx"
    s3.put_object(
        Bucket='files',
        Key=file_key,
        Body=docx_bytes,
        ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    cdn_url = f"https://cdn.poehali.dev/projects/{os.environ['AWS_ACCESS_KEY_ID']}/bucket/{file_key}"

    contract_no_esc = str(contract_no).replace("'", "''")
    conn2 = psycopg2.connect(dsn)
    cur2 = conn2.cursor()
    cur2.execute(
        f"INSERT INTO lead_documents (lead_id, doc_type, file_url, doc_no) "
        f"VALUES ({int(lead_id)}, 'contract', '{cdn_url}', '{contract_no_esc}') RETURNING id"
    )
    doc_id = cur2.fetchone()[0]
    conn2.commit()
    cur2.close()
    conn2.close()

    return {
        'statusCode': 200,
        'headers': {**cors_headers, 'Content-Type': 'application/json'},
        'body': json.dumps({'success': True, 'url': cdn_url, 'docId': doc_id}, ensure_ascii=False),
        'isBase64Encoded': False
    }